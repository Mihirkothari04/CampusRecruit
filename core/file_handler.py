import fitz  # pymupdf
import docx
import zipfile
import os
import tempfile

def extract_text_from_pdf(file_path_or_bytes):
    text = ""
    try:
        if isinstance(file_path_or_bytes, bytes):
            doc = fitz.open(stream=file_path_or_bytes, filetype="pdf")
        else:
            doc = fitz.open(file_path_or_bytes)
            
        for page in doc:
            text += page.get_text()
        doc.close()
    except Exception as e:
        print(f"Error reading PDF: {e}")
    return text

def extract_text_from_docx(file_path_or_bytes):
    text = ""
    try:
        if isinstance(file_path_or_bytes, bytes):
            import io
            doc = docx.Document(io.BytesIO(file_path_or_bytes))
        else:
            doc = docx.Document(file_path_or_bytes)
            
        for para in doc.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        print(f"Error reading DOCX: {e}")
    return text

def process_file_bytes(file_name, file_bytes):
    if file_name.lower().endswith('.pdf'):
        return extract_text_from_pdf(file_bytes)
    elif file_name.lower().endswith('.docx'):
        return extract_text_from_docx(file_bytes)
    return ""

def extract_from_zip(zip_bytes):
    import io
    extracted_files = []
    try:
        with zipfile.ZipFile(io.BytesIO(zip_bytes)) as z:
            for file_info in z.infolist():
                if file_info.is_dir():
                    continue
                if file_info.filename.lower().endswith(('.pdf', '.docx')) and not file_info.filename.startswith('__MACOSX'):
                    with z.open(file_info) as f:
                        content = f.read()
                        text = process_file_bytes(file_info.filename, content)
                        if text:
                            extracted_files.append({"file_name": os.path.basename(file_info.filename), "raw_text": text})
    except Exception as e:
        print(f"Error reading ZIP: {e}")
    return extracted_files
